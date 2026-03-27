"""DOCX export module in Lofte Studios format."""

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


async def create_docx(
    job_id: str,
    video_title: str,
    transcription: dict,
    translation: dict,
    ost_detection: dict,
    output_dir: str,
) -> str:
    """
    Create DOCX document in Lofte Studios format.

    Args:
        job_id: Job ID
        video_title: Video title
        transcription: Transcription data
        translation: Translation data
        ost_detection: OST detection data
        output_dir: Output directory

    Returns:
        Path to created DOCX file
    """
    logger.info(f"Creating DOCX for job {job_id}")

    try:
        output_path = Path(output_dir) / f"{job_id}_translation.docx"

        # Create document with landscape orientation
        doc = Document()

        # Set landscape orientation
        section = doc.sections[0]
        section.page_height = Inches(8.5)
        section.page_width = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Add header section
        _add_header(doc, video_title, transcription, ost_detection)

        # Add translation segments table
        _add_segments_table(doc, translation, ost_detection)

        # Save document
        doc.save(str(output_path))

        logger.info(f"Created DOCX: {output_path}")
        return str(output_path)

    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        raise RuntimeError(f"Failed to create DOCX: {str(e)}")


def _add_header(doc: Document, video_title: str, transcription: dict, ost_detection: dict):
    """Add header section with video metadata."""

    # Title
    title = doc.add_heading(video_title or "Video Translation", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    metadata_table = doc.add_table(rows=4, cols=2)
    metadata_table.style = "Light Grid Accent 1"

    # Duration
    duration = transcription.get("duration_seconds", 0)
    duration_str = _format_timestamp(duration)

    metadata_table.rows[0].cells[0].text = "Duration"
    metadata_table.rows[0].cells[1].text = duration_str

    metadata_table.rows[1].cells[0].text = "Segments"
    segments = transcription.get("segments", [])
    metadata_table.rows[1].cells[1].text = str(len(segments))

    metadata_table.rows[2].cells[0].text = "OST Items Detected"
    ost_items = ost_detection.get("ost_items", [])
    metadata_table.rows[2].cells[1].text = str(len(ost_items))

    # Video summary
    summary = doc.add_paragraph()
    summary.add_run("Summary: ").bold = True
    summary_text = metadata_table.rows[3].cells[0]

    doc.add_paragraph()  # Spacing


def _add_segments_table(doc: Document, translation: dict, ost_detection: dict):
    """Add main translation segments table."""

    segments = translation.get("segments", [])

    # Create 4-column table: Timestamp | Translation | OST | Notes
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Timestamp"
    header_cells[1].text = "Translation"
    header_cells[2].text = "OST"
    header_cells[3].text = "Notes"

    # Format header
    for cell in header_cells:
        _format_header_cell(cell)

    # Add segments
    for segment in segments:
        row = table.add_row()
        cells = row.cells

        # Timestamp (00.00.00 - 00.01.00 format with dots)
        timestamp_start = segment.get("timestamp_start", 0)
        timestamp_end = segment.get("timestamp_end", 0)
        timestamp_text = f"{_format_timestamp_with_dots(timestamp_start)} - {_format_timestamp_with_dots(timestamp_end)}"

        cells[0].text = timestamp_text
        _format_cell(cells[0], is_small=True)

        # Translation with speaker name bold
        speaker = segment.get("speaker")
        translation_text = segment.get("translation", "")

        paragraph = cells[1].paragraphs[0]
        if speaker:
            run = paragraph.add_run(f"{speaker}: ")
            run.bold = True
            paragraph.add_run(translation_text)
        else:
            paragraph.text = translation_text

        _format_cell(cells[1])

        # OST items
        ost_items = segment.get("ost_items", [])
        if ost_items:
            ost_text = " | ".join(
                [f"{item.get('type', 'unknown')}" for item in ost_items[:2]]
            )
            cells[2].text = ost_text
        else:
            cells[2].text = "-"

        _format_cell(cells[2], is_small=True)

        # Notes (term bold + explanation)
        notes = segment.get("notes", [])
        if notes:
            notes_para = cells[3].paragraphs[0]
            notes_para.clear()

            for note_idx, note in enumerate(notes):
                if note_idx > 0:
                    notes_para.add_run("\n")

                term = note.get("term", "")
                explanation = note.get("explanation", "")

                run = notes_para.add_run(f"{term}: ")
                run.bold = True
                notes_para.add_run(explanation)
        else:
            cells[3].text = "-"

        _format_cell(cells[3], is_small=True)


def _set_cell_shading(cell, color):
    """Set cell background shading color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _format_header_cell(cell):
    """Format header cell with blue-gray background and white text."""
    # Set background color to blue-gray (D9E2F3)
    _set_cell_shading(cell, 'D9E2F3')

    # Set text color and style
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # White


def _format_cell(cell, is_small: bool = False):
    """Format table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            if is_small:
                run.font.size = Pt(8)
            else:
                run.font.size = Pt(9)


def _format_timestamp(seconds: int) -> str:
    """Format timestamp as HH:MM:SS."""
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def _format_timestamp_with_dots(seconds: int) -> str:
    """Format timestamp as HH.MM.SS (with dots instead of colons)."""
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    return f"{hours:02d}.{minutes:02d}.{secs:02d}"
